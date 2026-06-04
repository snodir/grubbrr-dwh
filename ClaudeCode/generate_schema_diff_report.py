"""
GAS DB Schema Difference Report Generator
Produces: schema_diff_report_20260604.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────────────────────
C_GREEN  = RGBColor(0xD5, 0xE8, 0xD4)   # new / added
C_RED    = RGBColor(0xF8, 0xCE, 0xCC)   # removed
C_YELLOW = RGBColor(0xFF, 0xF2, 0xCC)   # changed
C_ORANGE = RGBColor(0xFF, 0xE6, 0xCC)   # breaking / high risk
C_BLUE   = RGBColor(0xDA, 0xE8, 0xFC)   # header rows
C_LGREY  = RGBColor(0xF5, 0xF5, 0xF5)   # alternating rows
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK   = RGBColor(0x1F, 0x4E, 0x79)   # deep blue for headings
C_MED    = RGBColor(0x2E, 0x75, 0xB6)   # medium blue

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_widths_inches):
    """Set column widths via tblGrid and cell widths."""
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        table._tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in col_widths_inches:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 1440)))
        tblGrid.append(gridCol)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(col_widths_inches[i] * 1440)))
            tcW.set(qn('w:type'), 'dxa')

def bold_para(para, text, size=11, color=None, align=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_DARK if level == 1 else C_MED
    return p

def add_table(doc, headers, rows, col_widths, row_colors=None):
    """
    row_colors: list of RGBColor per data row, or None for alternating grey/white.
    'H' as first element of a row means it's a sub-header row (blue bg, bold).
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], C_DARK)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(9)

    # Data rows
    for ri, row_data in enumerate(rows):
        is_subheader = (len(row_data) > 0 and row_data[0] == '__SUBHDR__')
        cells = table.rows[ri + 1].cells
        if is_subheader:
            for ci in range(n_cols):
                val = row_data[ci + 1] if ci + 1 < len(row_data) else ''
                set_cell_bg(cells[ci], C_BLUE)
                p = cells[ci].paragraphs[0]
                run = p.add_run(val)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = C_DARK
        else:
            if row_colors and ri < len(row_colors):
                bg = row_colors[ri]
            else:
                bg = C_LGREY if ri % 2 == 0 else C_WHITE
            for ci, val in enumerate(row_data):
                set_cell_bg(cells[ci], bg)
                p = cells[ci].paragraphs[0]
                p.add_run(str(val)).font.size = Pt(9)

    set_col_width(table, col_widths)
    doc.add_paragraph()
    return table

def risk_color(risk_str):
    r = risk_str.upper()
    if 'HIGH' in r or 'BREAKING' in r:
        return C_ORANGE
    if 'MEDIUM' in r:
        return C_YELLOW
    return C_GREEN

# ──────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('GAS DB Schema Difference Report')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = C_DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Test (Redesigned) vs Production')
r2.font.size = Pt(16)
r2.font.color.rgb = C_MED

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info_p.add_run('Pre-Promotion Review  |  Generated: 2026-06-04  |  DWH Team')
r3.font.size = Pt(11)
r3.italic = True

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = note_p.add_run('CONFIDENTIAL — Internal Use Only')
r4.bold = True
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
add_heading(doc, '1. Executive Summary')
p = doc.add_paragraph(
    'This report compares the Test (redesigned) and Production GAS DB PostgreSQL database schemas '
    'as of 2026-06-04, prior to promoting Test environment changes to Production. '
    'Both databases run PostgreSQL 16.9 on Ubuntu 20.04.'
)

doc.add_paragraph()
add_heading(doc, 'Key Findings', level=2)

findings = [
    ('1 new schema in Test', 'etl', C_GREEN),
    ('1 new type in Test', 'public.reportorigin (ENUM: Custom / System)', C_GREEN),
    ('9 tables only in Production', 'Candidates for archival / cleanup review', C_RED),
    ('49 tables only in Test', 'Net-new tables to be promoted', C_GREEN),
    ('32 tables in both with schema changes', 'Columns, types, indexes, or constraints differ', C_YELLOW),
    ('13 new sequences in Test', 'Auto-increment sequences for new/refactored tables', C_GREEN),
    ('41 new stored procedures in Test', 'New ETL silver-layer and dim refresh procedures', C_GREEN),
    ('1 function only in Production', 'fact.getsalesreport — verify before dropping', C_RED),
    ('3 function signatures changed', 'One high-risk core ETL function, two low-risk optimizations', C_YELLOW),
    ('2 breaking column type changes', 'dim.occasionsurvey.surveytype (text→int), stg...incidentid (text→bigint)', C_ORANGE),
]
add_table(doc,
    ['Finding', 'Detail'],
    [[f, d] for f, d, _ in findings],
    [3.0, 4.5],
    row_colors=[c for _, _, c in findings]
)

# ── 2. SCHEMA-LEVEL CHANGES ───────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '2. Schema-Level Changes')
add_table(doc,
    ['Schema', 'Status', 'Notes'],
    [
        ['dim',    'In Both',          'No structural changes'],
        ['fact',   'In Both',          'No structural changes'],
        ['ml',     'In Both',          'No structural changes'],
        ['public', 'In Both',          'Minor table changes (see §4.3)'],
        ['stg',    'In Both',          'Extensive new staging tables added (see §4.2)'],
        ['etl',    'TEST ONLY — NEW',  'New schema for ETL orchestration; 2 tables + 2 functions'],
    ],
    [1.5, 2.0, 4.0],
    row_colors=[C_WHITE, C_WHITE, C_WHITE, C_WHITE, C_WHITE, C_GREEN]
)

# ── 3. NEW TYPES ──────────────────────────────────────────────────────────────
add_heading(doc, '3. New Types / Enums')
add_table(doc,
    ['Type', 'Status', 'Details'],
    [['public.reportorigin', 'TEST ONLY — NEW', "ENUM with values: 'Custom', 'System'"]],
    [2.5, 1.5, 3.5],
    row_colors=[C_GREEN]
)

# ── 4. TABLES ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '4. Tables')

# 4.1 Prod-only
add_heading(doc, '4.1  Tables Only in Production  (9 tables — Review for Archival)', level=2)
doc.add_paragraph(
    'These tables exist in Production but NOT in Test. They may be legacy, temp, or deprecated objects. '
    'Do NOT drop them during promotion without explicit sign-off.'
)
add_table(doc,
    ['Schema', 'Table Name', 'Recommendation'],
    [
        ['fact', 'T_4442_1eca082546ae4881a2b72397e354c4fa', 'Review — likely temp / legacy hash table'],
        ['fact', 'T_5656_e6b4c2c790f549b8b4f741798c8912c6', 'Review — likely temp / legacy hash table'],
        ['fact', 'T_5731_5b365fb5f4a74dfeae4a5924eb4daf5b', 'Review — likely temp / legacy hash table'],
        ['fact', 'transactionheadertest',                  'Review — likely test artifact in Production'],
        ['fact', 'userbehaviour_reload',                   'Review — likely temp reload / migration table'],
        ['dim',  'locationgroup',                          'Review — removed in redesign; check FK refs'],
        ['dim',  'menuitem_bkp',                           'Review — backup table; archive or drop'],
        ['dim',  'menuitem_test',                          'Review — test artifact in Production'],
        ['dim',  'organizationparent',                     'Review — removed in redesign; check FK refs'],
    ],
    [0.8, 3.5, 3.2],
    row_colors=[C_RED] * 9
)

# 4.2 Test-only
add_heading(doc, '4.2  Tables Only in Test  (49 tables — To Be Promoted)', level=2)

rows_42 = [
    ['__SUBHDR__', 'dim schema — 4 new dimension tables', '', ''],
    ['dim', 'experiment',       'New experiment tracking dimension', C_GREEN],
    ['dim', 'holidays',         'New holidays reference dimension',  C_GREEN],
    ['dim', 'locationcatalog',  'New location catalog dimension',    C_GREEN],
    ['dim', 'menuentities',     'New menu entities dimension',       C_GREEN],
    ['__SUBHDR__', 'etl schema — 2 new ETL support tables', '', ''],
    ['etl', 'bronze_partition_registry',      'ETL partition tracking registry',    C_GREEN],
    ['etl', 'gas_db_object_dependency_sort',  'DB object dependency ordering',      C_GREEN],
    ['__SUBHDR__', 'fact schema — 4 new fact tables', '', ''],
    ['fact', 'devicehealth',                       'New device health fact table',              C_GREEN],
    ['fact', 'gem_failed_order_job_notifications', 'New GEM failed order notification fact',    C_GREEN],
    ['fact', 'recommendations_bkp',               'Recommendations backup table',              C_GREEN],
    ['fact', 'userbehaviour_exceptions',           'User behaviour exception tracking',         C_GREEN],
    ['__SUBHDR__', 'stg schema — 39 new staging tables (silver-layer ETL)', '', ''],
    ['stg', 'silver_transaction_header',           'Silver-layer transaction header staging',   C_GREEN],
    ['stg', 'silver_transaction_item',             'Silver-layer transaction item staging',     C_GREEN],
    ['stg', 'silver_transaction_payment',          'Silver-layer payment staging',              C_GREEN],
    ['stg', 'silver_transaction_refunds',          'Silver-layer refunds staging',              C_GREEN],
    ['stg', 'silver_kiosk_events',                 'Silver-layer kiosk events staging',         C_GREEN],
    ['stg', 'silver_item_modifiers',               'Silver-layer item modifiers staging',       C_GREEN],
    ['stg', 'silver_modifier_impressions',         'Silver-layer modifier impressions',         C_GREEN],
    ['stg', 'silver_modifier_interactions',        'Silver-layer modifier interactions',        C_GREEN],
    ['stg', 'silver_modifier_recommendations',     'Silver-layer modifier recommendations',     C_GREEN],
    ['stg', 'silver_upsell_recommendations',       'Silver-layer upsell recommendations',       C_GREEN],
    ['stg', 'silver_cep_incidents',                'Silver-layer CEP incidents',                C_GREEN],
    ['stg', 'silver_transaction_combo_items',      'Silver-layer combo items',                  C_GREEN],
    ['stg', 'lookup_silver_transaction_header',    'Lookup for silver tx header',               C_GREEN],
    ['stg', 'temp_silver_transaction_header',      'Temp staging for silver tx header',         C_GREEN],
    ['stg', 'dim_catalog',                         'Dim catalog staging',                       C_GREEN],
    ['stg', 'dim_category_hierarchy',              'Category hierarchy staging',                C_GREEN],
    ['stg', 'dim_cep_subscriptions',               'CEP subscriptions staging',                 C_GREEN],
    ['stg', 'dim_frequentcustomer',                'Frequent customer staging',                 C_GREEN],
    ['stg', 'dim_itemcategory',                    'Item category staging',                     C_GREEN],
    ['stg', 'dim_kiosk',                           'Kiosk staging',                             C_GREEN],
    ['stg', 'dim_kiosk_appearance',                'Kiosk appearance staging',                  C_GREEN],
    ['stg', 'dim_kiosk_config',                    'Kiosk config staging',                      C_GREEN],
    ['stg', 'dim_location_kiosks',                 'Location kiosks staging',                   C_GREEN],
    ['stg', 'dim_loyalty_configuration',           'Loyalty config staging',                    C_GREEN],
    ['stg', 'dim_menuitem',                        'Menu item staging',                         C_GREEN],
    ['stg', 'dim_modifier',                        'Modifier staging',                          C_GREEN],
    ['stg', 'dim_modifiergroup',                   'Modifier group staging',                    C_GREEN],
    ['stg', 'dim_modifiergroup_modifier_mapping',  'Modifier group mapping staging',            C_GREEN],
    ['stg', 'dim_occasionsurvey',                  'Occasion survey staging',                   C_GREEN],
    ['stg', 'dim_organization',                    'Organization staging',                      C_GREEN],
    ['stg', 'dim_organizationlocation',            'Org location staging',                      C_GREEN],
    ['stg', 'dim_payment_provider',                'Payment provider staging',                  C_GREEN],
    ['stg', 'dim_pos_provider',                    'POS provider staging',                      C_GREEN],
    ['stg', 'fact_devicestate',                    'Device state staging',                      C_GREEN],
    ['stg', 'fact_devicetelemetry',                'Device telemetry staging',                  C_GREEN],
    ['stg', 'fact_itemssurvey',                    'Items survey staging',                      C_GREEN],
    ['stg', 'fact_occasionsurveydetail',           'Occasion survey detail staging',            C_GREEN],
    ['stg', 'kioskdetails',                        'Kiosk details staging',                     C_GREEN],
]

# Build colour map per row
colors_42 = []
clean_rows_42 = []
for r in rows_42:
    if r[0] == '__SUBHDR__':
        clean_rows_42.append(['__SUBHDR__', r[1], '', ''])
        colors_42.append(C_BLUE)
    else:
        clean_rows_42.append(r[:3])
        colors_42.append(r[3])

add_table(doc,
    ['Schema', 'Table Name', 'Notes'],
    clean_rows_42,
    [0.7, 3.5, 3.3],
    row_colors=colors_42
)

# 4.3 Changed tables
doc.add_page_break()
add_heading(doc, '4.3  Tables in Both Environments — Schema Changes  (32 tables)', level=2)
doc.add_paragraph(
    'Color legend:  Green = column / index added   |   Red = column / index / constraint removed   |   '
    'Yellow = data type or nullability changed   |   Orange = breaking change'
)
doc.add_paragraph()

changed_tables = [
    {
        'name': 'dim.catalog',
        'rows': [
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.element',
        'rows': [
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.itemcategory  ★ Significant',
        'rows': [
            ['Column Added',     'sysinserttime',             'timestamp without time zone', C_GREEN],
            ['Column Added',     'sysupdatetime',             'timestamp without time zone', C_GREEN],
            ['Column Added',     'is_category_deleted',       'boolean',                     C_GREEN],
            ['Column Added',     'category_created_on',       'timestamp without time zone', C_GREEN],
            ['Column Added',     'category_modified_on',      'timestamp without time zone', C_GREEN],
            ['Column Added',     'is_alcoholic',              'boolean',                     C_GREEN],
            ['Column Added',     'number_of_items',           'smallint',                    C_GREEN],
            ['Column Added',     'number_of_sub_categories',  'smallint',                    C_GREEN],
            ['Column Added',     'number_of_item_variations', 'smallint',                    C_GREEN],
            ['Column Added',     'number_of_combos',          'smallint',                    C_GREEN],
            ['Column Added',     'number_of_combo_families',  'smallint',                    C_GREEN],
            ['Type Changed',     'isactive',                  'Prod: boolean DEFAULT true NOT NULL  →  Test: boolean nullable (no default)', C_YELLOW],
            ['Index Added',      'itemcategory_bkp_idx',      'UNIQUE on dim.itemcategory_bkp (locationid, categoryid)',                   C_GREEN],
            ['Index Added',      'itemcategory_bkp_locationid_idx', 'on dim.itemcategory_bkp (locationid) INCLUDE (categoryid, isactive)', C_GREEN],
        ]
    },
    {
        'name': 'dim.itemcategorymapping',
        'rows': [
            ['Column Added', 'locationid',   'text', C_GREEN],
            ['Column Added', 'categoryname', 'text', C_GREEN],
            ['Column Added', 'menuitemname', 'text', C_GREEN],
        ]
    },
    {
        'name': 'dim.kiosk',
        'rows': [
            ['Column Added',        'sysinserttime',           'timestamp without time zone',           C_GREEN],
            ['Column Added',        'sysupdatetime',           'timestamp without time zone',           C_GREEN],
            ['Constraint Removed',  'devicetype_not_empty',    "CHECK ((devicetype)::text <> ''::text) — removed in Test", C_RED],
            ['Index Removed',       'kioskid_idx',             'on dim.kiosk (locationid) INCLUDE (istestkiosk)',         C_RED],
        ]
    },
    {
        'name': 'dim.kioskdetails',
        'rows': [
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.location',
        'rows': [
            ['Column Added',  'sysinserttime',       'timestamp without time zone',                        C_GREEN],
            ['Column Added',  'sysupdatetime',       'timestamp without time zone',                        C_GREEN],
            ['Index Removed', 'locationgroupid_idx', 'on dim.location (locationgroupid)',                  C_RED],
            ['Index Removed', 'locationgroupid_uidx','on dim.locationgroup (locationgroupid)',             C_RED],
        ]
    },
    {
        'name': 'dim.occasionsurvey  ⚠ BREAKING',
        'rows': [
            ['Type Changed — BREAKING', 'surveytype',       'Prod: text  →  Test: integer — data migration required', C_ORANGE],
            ['Index Removed',           'organizationid_idx','on dim.occasionsurvey (organizationid)',                C_RED],
            ['Index Removed',           'surveyid_idx',      'on dim.occasionsurvey (surveyid) INCLUDE (surveyname, surveytype)', C_RED],
        ]
    },
    {
        'name': 'dim.ordertype',
        'rows': [
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.organization  ★ NOT NULL relaxed',
        'rows': [
            ['Column Added',        'sysinserttime',    'timestamp without time zone',           C_GREEN],
            ['Column Added',        'sysupdatetime',    'timestamp without time zone',           C_GREEN],
            ['Nullability Changed', 'organizationtype', 'Prod: smallint NOT NULL  →  Test: smallint nullable', C_YELLOW],
            ['Nullability Changed', 'status',           'Prod: smallint NOT NULL  →  Test: smallint nullable', C_YELLOW],
            ['Nullability Changed', 'createdon',        'Prod: timestamp NOT NULL  →  Test: timestamp nullable',C_YELLOW],
            ['Nullability Changed', 'active',           'Prod: boolean NOT NULL  →  Test: boolean nullable',   C_YELLOW],
        ]
    },
    {
        'name': 'dim.organizationlocation  ★ NOT NULL relaxed',
        'rows': [
            ['Column Added',        'sysinserttime',     'timestamp without time zone',           C_GREEN],
            ['Column Added',        'sysupdatetime',     'timestamp without time zone',           C_GREEN],
            ['Nullability Changed', 'organizationname',  'Prod: varchar(255) NOT NULL  →  Test: varchar(255) nullable', C_YELLOW],
            ['Nullability Changed', 'organizationtype',  'Prod: smallint NOT NULL  →  Test: smallint nullable',         C_YELLOW],
        ]
    },
    {
        'name': 'dim.upsellgrouplookup',
        'rows': [
            ['Column Added', 'catalogid',     'text',                        C_GREEN],
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.userlocation',
        'rows': [
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'dim.view  ★ Sequence added for PK',
        'rows': [
            ['Column Added',  'sysinserttime', 'timestamp without time zone',                          C_GREEN],
            ['Column Added',  'sysupdatetime', 'timestamp without time zone',                          C_GREEN],
            ['Type Changed',  'viewid',        "Prod: integer NOT NULL  →  Test: integer DEFAULT nextval('dim.view_id_seq') nullable", C_YELLOW],
            ['Index Removed', 'idx_view_viewid','on dim.view (viewid)',                                C_RED],
        ]
    },
    {
        'name': 'fact.deviceevent',
        'rows': [
            ['Index Added', 'ix_deviceevent_syscosmosts_brin', 'USING brin (syscosmosts) — performance improvement', C_GREEN],
        ]
    },
    {
        'name': 'fact.itemmodifier',
        'rows': [
            ['Nullability Changed', 'modifierprice',                   'Prod: numeric(12,3) NOT NULL  →  Test: numeric(12,3) nullable', C_YELLOW],
            ['Index Added',         'idx_fact_itemmodifier_locationid', 'on fact.itemmodifier (locationid)',                           C_GREEN],
        ]
    },
    {
        'name': 'fact.itemssurvey  ★ NOT NULL relaxed',
        'rows': [
            ['Nullability Changed', 'locationid',   'Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
            ['Nullability Changed', 'surveytransid','Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
            ['Nullability Changed', 'orderid',      'Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
            ['Nullability Changed', 'itemid',       'Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
        ]
    },
    {
        'name': 'fact.modifier_interactions  ★ NOT NULL relaxed',
        'rows': [
            ['Nullability Changed', 'locationid', 'Prod: text NOT NULL  →  Test: text nullable',  C_YELLOW],
            ['Nullability Changed', 'menuitemid', 'Prod: text NOT NULL  →  Test: text nullable',  C_YELLOW],
        ]
    },
    {
        'name': 'fact.occasionsurveydetail  ★ NOT NULL relaxed',
        'rows': [
            ['Nullability Changed', 'locationid', 'Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
            ['Nullability Changed', 'orderid',    'Prod: text NOT NULL  →  Test: text nullable', C_YELLOW],
        ]
    },
    {
        'name': 'fact.ordertiming',
        'rows': [
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'fact.pipelinerunstatus  ★ Index INCLUDE reduced',
        'rows': [
            ['Index Changed', 'idx_fact_pipelinerunstatus_correlationid',
             'Prod INCLUDES (pipelinestatus, pipelinetriggertime, pipelinecompletedtime)  →  Test INCLUDES (pipelinestatus) only. '
             'Queries that relied on covered pipelinetriggertime/pipelinecompletedtime will now require heap access.',
             C_YELLOW],
        ]
    },
    {
        'name': 'fact.transactionheader  ★ High-risk column change',
        'rows': [
            ['Nullability Changed', 'updateddate',
             'Prod: timestamp DEFAULT CURRENT_TIMESTAMP NOT NULL  →  Test: timestamp nullable, no default — confirm intentional',
             C_ORANGE],
            ['Index Added', 'ix_transactionheader_syscosmosts_brin', 'USING brin (syscosmosts)', C_GREEN],
        ]
    },
    {
        'name': 'fact.transactionitem  ★ Column order changed',
        'rows': [
            ['Column Reorder', '(multiple)', 'Column order differs between environments — verify no positional SELECT * dependencies', C_YELLOW],
            ['Index Added',    'idx_fact_transactionitem_locationid', 'on fact.transactionitem (locationid)',        C_GREEN],
            ['Index Removed',  'idx_transactionitem_headerid',        'on fact.transactionitem (transactionheaderid)', C_RED],
        ]
    },
    {
        'name': 'fact.transactionpayment',
        'rows': [
            ['Column Added',  'syscosmosts',                             'bigint',                                          C_GREEN],
            ['Index Added',   'idx_fact_transactionpayment_locationid',  'on fact.transactionpayment (locationid)',         C_GREEN],
            ['Index Removed', 'transactionpayment_orderid_idx',          'on fact.transactionpayment (orderid)',            C_RED],
        ]
    },
    {
        'name': 'fact.transactionrefunds',
        'rows': [
            ['Nullability Changed', 'locationid', 'Prod: varchar(50) NOT NULL  →  Test: varchar(50) nullable', C_YELLOW],
        ]
    },
    {
        'name': 'fact.userbehaviour',
        'rows': [
            ['Column Added', 'sysupdatetime',                    'timestamp without time zone',            C_GREEN],
            ['Index Added',  'ix_userbehaviour_syscosmosts_brin','USING brin (syscosmosts)',               C_GREEN],
        ]
    },
    {
        'name': 'fact.usercheckedin',
        'rows': [
            ['Column Added', 'syscosmosts', 'bigint', C_GREEN],
        ]
    },
    {
        'name': 'fact.vw_offer_analysis  ★ 6 new columns',
        'rows': [
            ['Column Added', 'offereditem_upselllevel',  'text', C_GREEN],
            ['Column Added', 'offered_promptitemid',     'text', C_GREEN],
            ['Column Added', 'offered_upsellgroupid',    'text', C_GREEN],
            ['Column Added', 'selecteditem_upselllevel', 'text', C_GREEN],
            ['Column Added', 'selected_promptitemid',    'text', C_GREEN],
            ['Column Added', 'selected_upsellgroupid',   'text', C_GREEN],
        ]
    },
    {
        'name': 'fact.watermarktable',
        'rows': [
            ['Column Added', 'sysinserttime', 'timestamp without time zone', C_GREEN],
            ['Column Added', 'sysupdatetime', 'timestamp without time zone', C_GREEN],
        ]
    },
    {
        'name': 'public.report  ★ Column reorder',
        'rows': [
            ['Column Reorder', 'origin', 'Prod: origin positioned before supportedformats  →  Test: origin after previewformat. '
             'Verify no positional column references.', C_YELLOW],
        ]
    },
    {
        'name': 'public.reportschedule  ★ NOT NULL relaxed',
        'rows': [
            ['Nullability Changed', 'periodfrom', 'Prod: timestamp NOT NULL  →  Test: timestamp nullable', C_YELLOW],
            ['Nullability Changed', 'periodto',   'Prod: timestamp NOT NULL  →  Test: timestamp nullable', C_YELLOW],
        ]
    },
    {
        'name': 'stg.gem_failed_order_job_notifications  ⚠ BREAKING',
        'rows': [
            ['Type Changed — BREAKING', 'incidentid', 'Prod: text  →  Test: bigint — data migration / cast required', C_ORANGE],
        ]
    },
]

for tbl_info in changed_tables:
    p = doc.add_paragraph()
    run = p.add_run(tbl_info['name'])
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_DARK

    add_table(doc,
        ['Change Type', 'Column / Object', 'Detail'],
        [r[:3] for r in tbl_info['rows']],
        [2.0, 2.0, 3.5],
        row_colors=[r[3] for r in tbl_info['rows']]
    )

# ── 5. SEQUENCES ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. Sequences')
add_heading(doc, '5.1  New Sequences in Test  (13 — to be promoted)', level=2)
seq_rows = [
    ['dim.element_elementid_seq',            'Auto-increment for dim.element'],
    ['dim.experiment_dimkey_seq',            'Auto-increment for dim.experiment'],
    ['dim.frequentcustomer_customerkey_seq', 'Auto-increment for dim.frequentcustomer'],
    ['dim.itemcategory_id_seq',              'Auto-increment for dim.itemcategory'],
    ['dim.kiosk_id_seq',                     'Auto-increment for dim.kiosk'],
    ['dim.menuitem_id_seq',                  'Auto-increment for dim.menuitem'],
    ['dim.occasionsurvey_surveykey_seq',     'Auto-increment for dim.occasionsurvey'],
    ['dim.ordertype_id_seq',                 'Auto-increment for dim.ordertype'],
    ['dim.view_id_seq',                      'Auto-increment for dim.view (also changes viewid DEFAULT — see §4.3)'],
    ['fact.devicestate_id_seq',              'Auto-increment for fact.devicestate'],
    ['fact.ordertiming_id_seq',              'Auto-increment for fact.ordertiming'],
    ['fact.transactionheader_id_seq',        'Auto-increment for fact.transactionheader'],
    ['fact.userbehaviour_id_seq',            'Auto-increment for fact.userbehaviour'],
]
add_table(doc,
    ['Sequence Name', 'Notes'],
    seq_rows,
    [3.5, 4.0],
    row_colors=[C_GREEN] * len(seq_rows)
)
doc.add_paragraph(
    'Action required: after promotion, set each sequence START value to max(existing_pk) + 1 '
    'in Production to avoid primary key conflicts.'
)

# ── 6. FUNCTIONS & STORED PROCEDURES ─────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6. Functions & Stored Procedures')

add_heading(doc, '6.1  Functions Only in Production  (1 — review before dropping)', level=2)
add_table(doc,
    ['Function', 'Recommendation'],
    [['fact.getsalesreport(IN _organizationid text, IN _startdate text, IN _enddate text, '
      'OUT transactioncount integer, OUT salestotal numeric, OUT avgtransaction numeric, '
      'OUT avgguest numeric, OUT avgguesttransaction numeric, OUT loyaltysales numeric, OUT loyaltypct numeric)',
      'Verify if still called by any application or BI layer before removing from Production.']],
    [5.5, 2.0],
    row_colors=[C_RED]
)

add_heading(doc, '6.2  Functions Only in Test  (41 — to be promoted)', level=2)

fn_rows = [
    ['__SUBHDR__', 'dim schema — 18 new stored procedures', ''],
    ['dim', 'usp_grubbrr_install_base_all_devices()',          C_GREEN],
    ['dim', 'usp_refresh_catalog()',                           C_GREEN],
    ['dim', 'usp_refresh_category_hierarchy()',                C_GREEN],
    ['dim', 'usp_refresh_dim_location_kiosk_details()',        C_GREEN],
    ['dim', 'usp_refresh_element()',                           C_GREEN],
    ['dim', 'usp_refresh_frequentcustomer()',                  C_GREEN],
    ['dim', 'usp_refresh_itemcategory()',                      C_GREEN],
    ['dim', 'usp_refresh_kiosk()',                             C_GREEN],
    ['dim', 'usp_refresh_location_kiosk_details()',            C_GREEN],
    ['dim', 'usp_refresh_menuitem()',                          C_GREEN],
    ['dim', 'usp_refresh_modifier()',                          C_GREEN],
    ['dim', 'usp_refresh_modifiergroup()',                     C_GREEN],
    ['dim', 'usp_refresh_modifiergroup_modifier_mapping()',    C_GREEN],
    ['dim', 'usp_refresh_occasionsurvey()',                    C_GREEN],
    ['dim', 'usp_refresh_ordertype()',                         C_GREEN],
    ['dim', 'usp_refresh_organization()',                      C_GREEN],
    ['dim', 'usp_refresh_organizationlocation()',              C_GREEN],
    ['dim', 'usp_refresh_view()',                              C_GREEN],
    ['__SUBHDR__', 'etl schema — 2 new ETL functions', ''],
    ['etl', 'truncate_silver_staging()',                       C_GREEN],
    ['etl', 'usp_sort_db_objects_by_dependency()',             C_GREEN],
    ['__SUBHDR__', 'fact schema — 21 new stored procedures', ''],
    ['fact', 'parse_iso_timestamp(ts_string text)',            C_GREEN],
    ['fact', 'usp_gem_ordertiming_to_fact_ordertiming()',      C_GREEN],
    ['fact', 'usp_gem_usercheckedin_to_fact_usercheckedin()',  C_GREEN],
    ['fact', 'usp_gsh_devicehealth_to_fact_devicestate()',     C_GREEN],
    ['fact', 'usp_gsh_devicetelemetry_to_fact_devicetelemetry()', C_GREEN],
    ['fact', 'usp_modifier_recommendation_analysis()',         C_GREEN],
    ['fact', 'usp_nge_update_itemssurvey()',                   C_GREEN],
    ['fact', 'usp_silver_aborted_orders_and_items_to_fact()',  C_GREEN],
    ['fact', 'usp_silver_cep_incidents_to_fact_cep_incidents()', C_GREEN],
    ['fact', 'usp_silver_item_modifiers_to_fact()',            C_GREEN],
    ['fact', 'usp_silver_kiosk_events_to_fact_cep_incidents()', C_GREEN],
    ['fact', 'usp_silver_kiosk_events_to_fact_deviceevent()',  C_GREEN],
    ['fact', 'usp_silver_kiosk_events_to_fact_userbehaviour()',C_GREEN],
    ['fact', 'usp_silver_modifier_impressions_to_fact()',      C_GREEN],
    ['fact', 'usp_silver_modifier_interactions_to_fact()',     C_GREEN],
    ['fact', 'usp_silver_modifier_recommendations_to_fact()',  C_GREEN],
    ['fact', 'usp_silver_transaction_item_to_fact()',          C_GREEN],
    ['fact', 'usp_silver_transaction_payment_to_fact()',       C_GREEN],
    ['fact', 'usp_silver_transaction_refunds_to_fact()',       C_GREEN],
    ['fact', 'usp_silver_upsell_recommendations_to_fact()',    C_GREEN],
    ['fact', 'usp_stg_gem_failed_order_job_notifications_to_fact()', C_GREEN],
    ['fact', 'usp_stg_occasionsurveydetail_to_fact()',         C_GREEN],
]

clean_fn = []
colors_fn = []
for r in fn_rows:
    if r[0] == '__SUBHDR__':
        clean_fn.append(['__SUBHDR__', r[1], ''])
        colors_fn.append(C_BLUE)
    else:
        clean_fn.append(r[:2])
        colors_fn.append(r[2])

add_table(doc, ['Schema', 'Function / Procedure Name'], clean_fn, [0.8, 6.7], row_colors=colors_fn)

add_heading(doc, '6.3  Function Signature / Implementation Changes', level=2)
add_table(doc,
    ['Function', 'Production', 'Test', 'Risk'],
    [
        ['dim.array_to_text(a jsonb)',
         'LANGUAGE plpgsql (mutable)',
         'LANGUAGE sql IMMUTABLE STRICT',
         'LOW — optimization only'],
        ['dim.is_valid_jsonb(input text)',
         'LANGUAGE plpgsql IMMUTABLE',
         'LANGUAGE plpgsql IMMUTABLE STRICT (added STRICT)',
         'LOW — stricter null handling'],
        ['fact.usp_modifier_recommendations_stage_to_fact()',
         'Exists in Production',
         'Replaced by usp_silver_modifier_recommendations_to_fact()',
         'MEDIUM — verify ADF pipeline SP references'],
        ['fact.usp_silver_transaction_header_to_fact()',
         'Exists in both environments',
         'Body updated with more extensive silver-layer ETL logic',
         'HIGH — core ETL; run full regression'],
    ],
    [2.5, 1.8, 2.2, 1.0],
    row_colors=[C_GREEN, C_GREEN, C_YELLOW, C_ORANGE]
)

# ── 7. RISK ASSESSMENT ────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '7. Risk Assessment & Pre-Promotion Checklist')

add_heading(doc, '7.1  Breaking / High-Risk Changes', level=2)
add_table(doc,
    ['#', 'Object', 'Change', 'Risk', 'Action Required'],
    [
        ['1', 'dim.occasionsurvey.surveytype',
         'text → integer',
         'HIGH',
         'Data migration script needed; all existing text values must be numeric-castable'],
        ['2', 'stg.gem_failed_order_job_notifications.incidentid',
         'text → bigint',
         'HIGH',
         'Validate no non-numeric text values exist in Production before promotion'],
        ['3', 'fact.transactionheader.updateddate',
         'NOT NULL DEFAULT CURRENT_TIMESTAMP → nullable, no default',
         'HIGH',
         'Dropping NOT NULL drops data integrity guarantee; confirm intentional with data team'],
        ['4', 'fact.usp_silver_transaction_header_to_fact()',
         'SP body changed — new silver-layer ETL logic',
         'HIGH',
         'Run full end-to-end ETL regression test before promotion'],
        ['5', 'fact.usp_modifier_recommendations_stage_to_fact()',
         'Removed; replaced with new SP name',
         'MEDIUM',
         'Verify no ADF pipeline activity or scheduled job still references the old SP name'],
        ['6', 'dim.organization',
         '4 columns relaxed from NOT NULL to nullable',
         'MEDIUM',
         'Confirm intentional; audit downstream ETL/BI queries that assume NOT NULL'],
        ['7', 'dim.organizationlocation',
         '2 columns relaxed from NOT NULL to nullable',
         'MEDIUM',
         'Same as above'],
    ],
    [0.3, 2.0, 2.0, 0.8, 2.4],
    row_colors=[C_ORANGE, C_ORANGE, C_ORANGE, C_ORANGE, C_YELLOW, C_YELLOW, C_YELLOW]
)

add_heading(doc, '7.2  Medium-Risk Changes', level=2)
add_table(doc,
    ['#', 'Object', 'Change', 'Action Required'],
    [
        ['1', 'dim.kiosk',                   'CHECK constraint devicetype_not_empty removed', 'Confirm intentional; empty devicetype strings will now be allowed'],
        ['2', 'fact.pipelinerunstatus index', 'INCLUDE columns reduced in covering index',    'Benchmark queries that filter on pipelinetriggertime / pipelinecompletedtime'],
        ['3', 'fact.transactionitem',         'Column order changed',                         'Audit for SELECT * or positional column access in application / ETL code'],
        ['4', 'public.report',               'Column order changed (origin column)',          'Audit for SELECT * or positional column access'],
        ['5', 'dim.view.viewid',              'DEFAULT now uses sequence',                    'Set sequence start value = max(viewid) + 1 in Production post-promotion'],
        ['6', 'Multiple tables',             'Indexes removed (locationgroupid, surveyid, transactionitem)', 'Run EXPLAIN ANALYZE on key queries before and after in a staging environment'],
    ],
    [0.3, 2.0, 2.5, 2.7],
    row_colors=[C_YELLOW] * 6
)

add_heading(doc, '7.3  Low-Risk / Safe Changes  (Additive Only)', level=2)
safe_rows = [
    'Adding sysinserttime / sysupdatetime columns — all nullable, fully backward-compatible',
    'Adding syscosmosts bigint columns — nullable, additive',
    'All 49 new staging/fact/dim/etl tables — net-new, no existing dependencies',
    'All 13 new sequences — new objects, no conflict with existing data unless start values are wrong',
    'All 41 new stored procedures — new objects, no replacement risk',
    'New etl schema and all objects within it',
    'BRIN indexes on fact.transactionheader, fact.userbehaviour, fact.deviceevent — additive, performance only',
    'New locationid indexes on fact tables — additive, performance only',
    'dim.array_to_text and dim.is_valid_jsonb language/volatility optimization — functionally equivalent',
    '6 new columns in fact.vw_offer_analysis — additive, non-breaking',
]
for s in safe_rows:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(s).font.size = Pt(9)

add_heading(doc, '7.4  Promotion Checklist', level=2)
checklist = [
    ['1',  'Backup Production DB completely before any DDL changes',                          'DBA',          '[ ]'],
    ['2',  'Write and test data migration for dim.occasionsurvey.surveytype (text→int)',       'DWH Team',     '[ ]'],
    ['3',  'Validate stg.gem_failed_order_job_notifications.incidentid for non-numeric text', 'DWH Team',     '[ ]'],
    ['4',  'Confirm fact.transactionheader.updateddate nullable change is intentional',        'DWH Team',     '[ ]'],
    ['5',  'Run end-to-end ETL pipeline regression test in Test environment',                  'DWH Team',     '[ ]'],
    ['6',  'Update all ADF pipeline SP references from old names to new names',                'DWH/ADF Team', '[ ]'],
    ['7',  'Confirm 9 Prod-only tables: archive, keep, or drop plan with sign-off',            'DWH Team',     '[ ]'],
    ['8',  'Set all 13 sequence start values = max(pk) + 1 after promotion',                   'DBA',          '[ ]'],
    ['9',  'Run query performance tests after index removals / changes',                       'DWH Team',     '[ ]'],
    ['10', 'Validate all 41 new stored procedures are wired into ADF pipeline activities',     'ADF/DWH Team', '[ ]'],
    ['11', 'Review dim.organization & dim.organizationlocation NOT NULL relaxation with data team','DWH Team', '[ ]'],
    ['12', 'Audit SELECT * usage against tables with column reorder (transactionitem, report)','DWH Team',     '[ ]'],
    ['13', 'Smoke-test Production after promotion; confirm ADF pipeline runs successfully',    'DWH/ADF Team', '[ ]'],
]
add_table(doc,
    ['#', 'Item', 'Owner', 'Done'],
    checklist,
    [0.3, 5.5, 1.2, 0.5]
)

# ── 8. APPENDIX — VIEWS ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '8. Appendix — Views  (Unchanged in Both Environments)')
add_table(doc,
    ['View', 'Schema', 'Status'],
    [
        ['dim.businessdate',         'dim', 'No changes detected'],
        ['dim.vw_weatherhourlydata', 'dim', 'No changes detected'],
        ['dim.vworganizationlocation','dim','No changes detected'],
    ],
    [3.5, 1.0, 3.0]
)

# ── FOOTER NOTE ───────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Confidential — Internal Use Only  |  Grubbrr DWH Team  |  2026-06-04')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\user\Work\Grubbrr\grubbrr-dwh\ClaudeCode\schema_diff_report_20260604.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
